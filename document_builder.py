from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_marketing_strategic_document():
    doc = Document()
    
    # إعداد التنسيق العام والخطوط
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    # العنوان الرئيسي للمستند
    title = doc.add_heading('Marketing Analytics: Sentiment & ROI Strategic Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. الملخص التنفيذي
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This report provides a comprehensive analysis of marketing performance by integrating "
        "financial metrics (ROAS) with customer psychological states (Sentiment). The objective "
        "is to identify the most effective channels for sustainable growth and brand loyalty."
    )

    # 2. المنهجية المتبعة
    doc.add_heading('2. Analytical Methodology', level=1)
    doc.add_paragraph(
        "Data was processed from 300 customer interactions. Sentiment was extracted using a "
        "proprietary keyword-based NLP model, while ROI was calculated based on the ratio "
        "between total purchase value and advertising expenditure."
    )

    # 3. التحليل المرئي للبيانات
    doc.add_heading('3. Visual Data Insights', level=1)
    
    # القسم الأول: كفاءة القنوات
    doc.add_heading('3.1 Return on Ad Spend (ROAS) per Channel', level=2)
    if os.path.exists('figures/channel_roas.png'):
        doc.add_picture('figures/channel_roas.png', width=Inches(5.5))
        doc.add_paragraph(
            "Figure 1: This chart highlights the financial efficiency of each marketing channel, "
            "allowing for data-driven budget reallocation."
        )

    doc.add_page_break()

    # القسم الثاني: تحليل المشاعر
    doc.add_heading('3.2 Customer Sentiment Breakdown', level=2)
    if os.path.exists('figures/sentiment_breakdown.png'):
        doc.add_picture('figures/sentiment_breakdown.png', width=Inches(5.5))
        doc.add_paragraph(
            "Figure 2: A breakdown of customer emotions per channel. This metric identifies "
            "potential brand risks where high sales might coincide with negative feedback."
        )

    # 4. التوصيات الاستراتيجية والنمو
    doc.add_heading('4. Strategic Recommendations', level=1)
    doc.add_paragraph(
        "1. Scale budgets for channels showing high Positive Sentiment and ROAS.\n"
        "2. Audit the 'Negative' sentiment clusters to improve product descriptions or support.\n"
        "3. Transition towards a humanized marketing model that prioritizes long-term brand equity."
    )

    output_path = 'outputs/Marketing_Sentiment_Strategic_Report.docx'
    doc.save(output_path)
    print(f"Marketing Strategic Word document created at: {output_path}")

if __name__ == "__main__":
    create_marketing_strategic_document()